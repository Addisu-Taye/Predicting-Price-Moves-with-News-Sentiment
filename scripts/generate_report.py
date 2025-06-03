import pandas as pd
import matplotlib.pyplot as plt
from textblob import TextBlob
from docx import Document
from docx.shared import Inches
import os

# === Step 1: Load and Prepare Data ===

# File paths
news_file = "correlations/TSLA_data.csv"
price_file = "correlations/TSLA_historical_data.csv"

# Extract ticker from file name
ticker = os.path.basename(price_file).split("_")[0].upper()

# Load data
news_df = pd.read_csv(news_file, parse_dates=["date"])
price_df = pd.read_csv(price_file, parse_dates=["Date"])

# Normalize column names
news_df.columns = news_df.columns.str.lower()
price_df.columns = price_df.columns.str.lower()

# Add ticker column to price_df
price_df['stock'] = ticker

# Standardize columns
price_df.rename(columns={"date": "price_date", "close": "close_price"}, inplace=True)
price_df = price_df[["price_date", "stock", "close_price"]]
price_df.sort_values(by=["stock", "price_date"], inplace=True)

# === Step 2: Sentiment Analysis ===

def get_sentiment(text):
    return TextBlob(str(text)).sentiment.polarity

news_df['stock'] = news_df['stock'].str.upper()
news_df['sentiment'] = news_df['headline'].apply(get_sentiment)

# Daily average sentiment per stock
sentiment_daily = news_df.groupby(['date', 'stock'])['sentiment'].mean().reset_index()

# === Step 3: Daily Return Calculation ===

price_df['daily_return'] = price_df.groupby('stock')['close_price'].pct_change()

# === Step 4: Merge Sentiment and Price ===

merged = pd.merge(
    sentiment_daily,
    price_df,
    left_on=["date", "stock"],
    right_on=["price_date", "stock"],
    how="inner"
)

# === Step 5: Correlation Analysis with Safe Handling ===

# Drop rows with missing sentiment or return
valid_merged = merged.dropna(subset=["sentiment", "daily_return"])

# Filter and compute correlation for each stock as a Series
filtered = valid_merged.groupby("stock").filter(lambda x: len(x.dropna()) > 2)
correlation = (
    filtered.groupby("stock")[["sentiment", "daily_return"]]
    .apply(lambda x: x["sentiment"].corr(x["daily_return"]))
)

correlation = correlation.dropna().sort_values(ascending=False)


# Handle plotting
if correlation.empty:
    print("⚠️ No valid correlation data to plot. Skipping chart.")
    correlation_plot_file = None
else:
    plt.figure(figsize=(10, 5))
    correlation.plot(kind='bar', title="Sentiment vs Daily Return Correlation per Stock")
    plt.ylabel("Pearson Correlation")
    plt.tight_layout()
    plt.savefig("correlation_plot.png")
    plt.close()
    correlation_plot_file = "correlation_plot.png"

# === Step 6: EDA for Report ===

headline_lengths = news_df["headline"].str.len()
top_publishers = news_df["publisher"].value_counts().head(5)

# === Step 7: Generate Word Report ===

doc = Document()

doc.add_heading("Week 1 Assignment Report", 0)
doc.add_heading("Predicting Price Moves with News Sentiment", level=1)

doc.add_paragraph(
    "This report analyzes financial news headlines and their correlation with stock price movements. "
    "Using sentiment analysis and historical stock data, we explore whether sentiment can help predict stock returns."
)

# Dataset Overview
doc.add_heading("1. Dataset Overview", level=1)
doc.add_paragraph(f"Total headlines: {len(news_df)}")
doc.add_paragraph(f"Total stocks in price data: {price_df['stock'].nunique()}")

# Task 1 - EDA
doc.add_heading("2. Exploratory Data Analysis", level=1)
doc.add_paragraph("Top 5 publishers by article count:")
for pub, count in top_publishers.items():
    doc.add_paragraph(f"{pub}: {count} articles")

doc.add_paragraph("Headline length stats:")
doc.add_paragraph(str(headline_lengths.describe()))

# Task 2 - Technical Analysis
doc.add_heading("3. Technical Analysis", level=1)
doc.add_paragraph("Daily stock returns were computed as the percentage change in the closing price.")

# Task 3 - Sentiment Correlation
doc.add_heading("4. Correlation Analysis", level=1)
doc.add_paragraph(
    "Average daily sentiment scores were computed from financial news headlines. "
    "These were then matched to corresponding daily stock returns to compute the Pearson correlation."
)

# Correlation table
if not correlation.empty:
    doc.add_paragraph("Sentiment vs. Return Correlation:")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Stock"
    hdr_cells[1].text = "Correlation"
    for stock, corr in correlation.items():
        row = table.add_row().cells
        row[0].text = stock
        row[1].text = f"{corr:.4f}"
else:
    doc.add_paragraph("No sufficient data to calculate sentiment-return correlation.")

if correlation_plot_file:
    doc.add_picture(correlation_plot_file, width=Inches(6))

# Final observations
doc.add_heading("5. Observations and Recommendations", level=1)
doc.add_paragraph(
    "Some tickers may show stronger correlation between sentiment and returns, indicating predictive potential. "
    "In cases where data was insufficient, we recommend improving date alignment and ensuring richer headline coverage."
)

# Save report
output_file = "Week1_Final_Report.docx"
doc.save(output_file)

# Cleanup
if correlation_plot_file:
    os.remove(correlation_plot_file)

print(f"✅ Report successfully generated: {output_file}")
